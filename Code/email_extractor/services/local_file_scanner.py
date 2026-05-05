"""
services/local_file_scanner.py
Рекурсивное сканирование локальной директории для извлечения email-адресов.

Реализует ILocalFileScanner.
Поддерживает: txt, csv, xlsx/xls, zip/tar/gz/bz2/rar/7z и любые текстовые файлы.
"""
from __future__ import annotations

import csv
import io
import logging
import os
from pathlib import Path

from ..core.entities import Contact
from ..core.interfaces import ILocalFileScanner
from ..services.email_extractor import EMAIL_RE, is_fake_email

log = logging.getLogger(__name__)


class LocalFileScanner(ILocalFileScanner):
    """
    Рекурсивно обходит директорию и извлекает email-адреса из всех файлов.

    Поддерживаемые форматы:
      - .xlsx / .xls  — через pandas (если установлен)
      - .csv          — умный CSV-парсер
      - .zip .tar .gz .bz2 .rar .7z — через patoolib (если установлен)
      - всё остальное — попытка прочитать как текст

    Args:
        extensions: Множество расширений для фильтрации (напр. ``{'.csv', '.txt'}``).
                    Передайте ``{'*'}`` для обхода всех файлов.
    """

    def __init__(self, extensions: set[str] | None = None) -> None:
        self._extensions: set[str] = extensions or {"*"}

    # ------------------------------------------------------------------
    # ILocalFileScanner implementation
    # ------------------------------------------------------------------

    def scan(self, directory: Path) -> list[Contact]:
        """Рекурсивно обойти *directory* и вернуть все найденные контакты."""
        if not directory.exists():
            log.warning("📂 Директория не существует: %s", directory)
            return []

        files = self._find_files(directory)
        contacts: dict[str, Contact] = {}

        for filepath in files:
            try:
                new = self._extract_from_file(filepath)
                for c in new:
                    if c.email not in contacts:
                        contacts[c.email] = c
            except Exception as exc:
                log.debug("✗ Ошибка обработки %s: %s", filepath.name, exc)

        log.info("📂 Локальное сканирование: %d файлов → %d контактов", len(files), len(contacts))
        return list(contacts.values())

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _find_files(self, directory: Path) -> list[Path]:
        """Рекурсивно найти все файлы с нужными расширениями."""
        found: list[Path] = []
        try:
            for entry in directory.iterdir():
                if entry.is_file():
                    if self._extensions == {"*"} or entry.suffix.lower() in self._extensions:
                        found.append(entry)
                elif entry.is_dir():
                    found.extend(self._find_files(entry))
        except (PermissionError, OSError) as exc:
            log.debug("✗ Нет доступа к %s: %s", directory, exc)
        return found

    def _extract_from_file(self, filepath: Path) -> list[Contact]:
        """Выбрать стратегию парсинга по расширению файла."""
        ext = filepath.suffix.lower()
        if ext in (".xlsx", ".xls"):
            return self._from_excel(filepath)
        elif ext == ".csv":
            return self._from_csv(filepath)
        elif ext == ".pdf":
            return self._from_pdf(filepath)
        elif ext in (".docx", ".doc"):
            return self._from_docx(filepath)
        elif ext in (".zip", ".tar", ".gz", ".bz2", ".rar", ".7z"):
            return self._from_archive(filepath)
        else:
            return self._from_text(filepath)

    # ------------------------------------------------------------------
    # Format-specific parsers
    # ------------------------------------------------------------------

    @staticmethod
    def _from_text(filepath: Path) -> list[Contact]:
        """Прочитать файл построчно как текст и найти email через regex."""
        seen: dict[str, Contact] = {}
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    for email in EMAIL_RE.findall(line):
                        e = email.lower().strip()
                        if not is_fake_email(e) and e not in seen:
                            seen[e] = Contact.from_email_only(e)
        except Exception as exc:
            log.debug("✗ Ошибка чтения текста %s: %s", filepath.name, exc)
        return list(seen.values())

    @staticmethod
    def _from_excel(filepath: Path) -> list[Contact]:
        """Читать Excel через pandas (опциональная зависимость)."""
        try:
            import pandas as pd  # type: ignore[import]

            df_dict = pd.read_excel(filepath, sheet_name=None, header=None)
            text = "\n".join(
                df.to_string(index=False, header=False) for df in df_dict.values()
            )
            seen: dict[str, Contact] = {}
            for e in EMAIL_RE.findall(text):
                el = e.lower().strip()
                if not is_fake_email(el) and el not in seen:
                    seen[el] = Contact.from_email_only(el)
            return list(seen.values())
        except Exception as exc:
            log.debug("✗ Ошибка Excel %s: %s", filepath.name, exc)
            return []

    @staticmethod
    def _from_csv(filepath: Path) -> list[Contact]:
        """Умный CSV-парсер с построчным чтением."""
        seen: dict[str, Contact] = {}
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.reader(f)
                rows = []
                
                for row_idx, row in enumerate(reader):
                    # Попутно проверяем регуляркой все ячейки (надежно)
                    for cell in row:
                        for e in EMAIL_RE.findall(cell):
                            el = e.lower().strip()
                            if not is_fake_email(el) and el not in seen:
                                seen[el] = Contact.from_email_only(el)
        except Exception as exc:
            log.debug("✗ Ошибка CSV %s: %s", filepath.name, exc)
            
        return list(seen.values())

    @staticmethod
    def _from_pdf(filepath: Path) -> list[Contact]:
        """Извлечь email из PDF через pdfplumber (опциональная зависимость)."""
        try:
            import pdfplumber  # type: ignore[import]
        except ImportError:
            log.debug("pdfplumber не установлен — PDF %s пропущен", filepath.name)
            return []

        seen: dict[str, Contact] = {}
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if not text:
                        continue
                    for email in EMAIL_RE.findall(text):
                        e = email.lower().strip()
                        if not is_fake_email(e) and e not in seen:
                            seen[e] = Contact.from_email_only(e)
        except Exception as exc:
            log.debug("✗ Ошибка PDF %s: %s", filepath.name, exc)
        return list(seen.values())

    @staticmethod
    def _from_docx(filepath: Path) -> list[Contact]:
        """Извлечь email из DOCX через python-docx (опциональная зависимость)."""
        try:
            import docx  # type: ignore[import]
        except ImportError:
            log.debug("python-docx не установлен — DOCX %s пропущен", filepath.name)
            return []

        seen: dict[str, Contact] = {}
        try:
            doc = docx.Document(filepath)
            for paragraph in doc.paragraphs:
                for email in EMAIL_RE.findall(paragraph.text):
                    e = email.lower().strip()
                    if not is_fake_email(e) and e not in seen:
                        seen[e] = Contact.from_email_only(e)
            # Также проверяем таблицы внутри документа
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for email in EMAIL_RE.findall(cell.text):
                            e = email.lower().strip()
                            if not is_fake_email(e) and e not in seen:
                                seen[e] = Contact.from_email_only(e)
        except Exception as exc:
            log.debug("✗ Ошибка DOCX %s: %s", filepath.name, exc)
        return list(seen.values())

    def _from_archive(self, filepath: Path) -> list[Contact]:
        """Распаковать архив во временную директорию и рекурсивно обойти."""
        import shutil
        import tempfile

        try:
            import patoolib  # type: ignore[import]
        except ImportError:
            log.debug("patoolib не установлен — архив %s пропущен", filepath.name)
            return []

        temp_dir = tempfile.mkdtemp()
        try:
            patoolib.extract_archive(str(filepath), outdir=temp_dir, interactive=False)
            return self.scan(Path(temp_dir))
        except Exception as exc:
            log.debug("✗ Ошибка распаковки %s: %s", filepath.name, exc)
            return []
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)
